from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import fitz
from PySide6.QtCore import QRectF, Qt
from PySide6.QtGui import QFont, QImage, QPainter, QTextDocument, QTextOption
from PySide6.QtWidgets import QFileDialog, QWidget


@dataclass
class LoadedDocument:
    path: Path
    doc: fitz.Document | None
    source_type: str
    fallback_pages: list[QImage] | None = None


class DocumentLoader:
    def __init__(self) -> None:
        self.loaded_documents: list[LoadedDocument] = []
        self.doc_index: int = -1
        self.page_index: int = 0
        self._temp_dir = Path(tempfile.mkdtemp(prefix='doc_capture_source_'))
        try:
            fitz.TOOLS.mupdf_display_warnings(False)
        except Exception:
            pass
        try:
            fitz.TOOLS.mupdf_display_errors(False)
        except Exception:
            pass

    def has_document(self) -> bool:
        return 0 <= self.doc_index < len(self.loaded_documents)

    def current_document(self) -> Optional[LoadedDocument]:
        if not self.has_document():
            return None
        return self.loaded_documents[self.doc_index]

    def open_file_dialog(self, parent: QWidget) -> bool:
        paths, _ = QFileDialog.getOpenFileNames(
            parent,
            '문서 불러오기',
            '',
            '지원 문서 (*.pdf *.doc *.docx *.hwp *.hwpx);;PDF Files (*.pdf);;Word Files (*.doc *.docx);;Hancom Files (*.hwp *.hwpx)',
        )
        if not paths:
            return False
        loaded_any = False
        for path in paths:
            loaded_any = self.open_document(path) or loaded_any
        return loaded_any

    def open_document(self, path: str) -> bool:
        src = Path(path)
        ext = src.suffix.lower()
        if ext == '.pdf':
            loaded = LoadedDocument(src, fitz.open(str(src)), 'pdf')
        elif ext in {'.doc', '.docx'}:
            loaded = self._open_word_family(src)
        elif ext in {'.hwp', '.hwpx'}:
            loaded = self._open_hwp_family(src)
        else:
            raise ValueError(f'지원하지 않는 형식입니다: {ext}')

        self.loaded_documents.append(loaded)
        self.doc_index = len(self.loaded_documents) - 1
        self.page_index = 0
        return True

    def _open_word_family(self, src: Path) -> LoadedDocument:
        pdf_doc = self._open_word_via_pdf_bridge(src)
        if pdf_doc is None:
            pdf_doc = self._open_word_via_libreoffice_pdf_bridge(src)
        if pdf_doc is not None:
            return LoadedDocument(src, pdf_doc, src.suffix.lower().lstrip('.'))
        text = ''
        if src.suffix.lower() == '.docx':
            text = self._extract_docx_text(src)
        else:
            text = self._extract_doc_text(src)
        if not text.strip():
            if src.suffix.lower() == '.doc':
                text = self._build_unreadable_doc_notice(src)
            else:
                raise RuntimeError('DOC/DOCX에서 표시 가능한 내용을 읽지 못했습니다. Word COM PDF 변환 또는 읽기 전용 preview 경로 모두 실패했습니다.')
        return LoadedDocument(src, None, src.suffix.lower().lstrip('.'), fallback_pages=self._text_to_pages(text, title=src.name))


    def _open_word_via_pdf_bridge(self, src: Path) -> fitz.Document | None:
        try:
            import pythoncom  # type: ignore
            import win32com.client  # type: ignore
        except Exception:
            return None
        word = None
        doc = None
        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(str(src), ReadOnly=True)
            out_pdf = self._temp_dir / f'{src.stem}_{len(self.loaded_documents)+1}.pdf'
            doc.ExportAsFixedFormat(str(out_pdf), 17)
            if out_pdf.exists():
                return fitz.open(str(out_pdf))
        except Exception:
            return None
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
        return None

    def _open_word_via_libreoffice_pdf_bridge(self, src: Path) -> fitz.Document | None:
        soffice = shutil.which('soffice')
        if soffice is None:
            return None
        out_dir = self._temp_dir / f'libreoffice_pdf_{len(self.loaded_documents) + 1}'
        out_dir.mkdir(parents=True, exist_ok=True)
        try:
            result = subprocess.run(
                [soffice, '--headless', '--convert-to', 'pdf', '--outdir', str(out_dir), str(src)],
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )
        except Exception:
            return None
        if result.returncode != 0:
            return None
        out_pdf = out_dir / f'{src.stem}.pdf'
        if not out_pdf.exists():
            candidates = sorted(out_dir.glob('*.pdf'))
            if not candidates:
                return None
            out_pdf = candidates[0]
        try:
            return fitz.open(str(out_pdf))
        except Exception:
            return None

    def _open_hwp_family(self, src: Path) -> LoadedDocument:
        pdf_doc = self._open_hwp_via_pdf_bridge(src)
        if pdf_doc is None:
            pdf_doc = self._open_hwp_via_pyhwpx_bridge(src)
        if pdf_doc is not None:
            return LoadedDocument(src, pdf_doc, src.suffix.lower().lstrip('.'))
        fallback_pages = self._render_hwp_text_fallback(src)
        if fallback_pages:
            return LoadedDocument(src, None, src.suffix.lower().lstrip('.'), fallback_pages=fallback_pages)
        raise RuntimeError('HWP/HWPX를 읽지 못했습니다. simple-hwp2pdf, pyhwpx(한/글 설치), helper-hwp/pyhwp 계열 fallback도 모두 실패했습니다.')


    def _open_hwp_via_pyhwpx_bridge(self, src: Path) -> fitz.Document | None:
        try:
            from pyhwpx import Hwp  # type: ignore
        except Exception:
            return None
        hwp = None
        out_pdf = self._temp_dir / f'{src.stem}_{len(self.loaded_documents)+1}_pyhwpx.pdf'
        try:
            hwp = Hwp()
            if hasattr(hwp, 'open'):
                hwp.open(str(src))
            elif hasattr(hwp, 'Open'):
                hwp.Open(str(src))
            else:
                return None
            if hasattr(hwp, 'save_pdf_as_image'):
                hwp.save_pdf_as_image(str(out_pdf))
            elif hasattr(hwp, 'save_as'):
                try:
                    hwp.save_as(path=str(out_pdf), format='PDF')
                except TypeError:
                    hwp.save_as(str(out_pdf), format='PDF')
            elif hasattr(hwp, 'SaveAs'):
                hwp.SaveAs(str(out_pdf), 'PDF')
            if out_pdf.exists():
                return fitz.open(str(out_pdf))
        except Exception:
            return None
        finally:
            for name in ('quit', 'Quit', 'close', 'Close'):
                try:
                    if hwp is not None and hasattr(hwp, name):
                        getattr(hwp, name)()
                        break
                except Exception:
                    pass
        return None

    def _open_hwp_via_pdf_bridge(self, src: Path) -> fitz.Document | None:
        try:
            from simple_hwp2pdf import convert  # type: ignore
            out_pdf = self._temp_dir / f'{src.stem}_{len(self.loaded_documents)+1}.pdf'
            method = 'standalone' if src.suffix.lower() == '.hwpx' else 'auto'
            convert(str(src), str(out_pdf), method=method)
            if out_pdf.exists():
                return fitz.open(str(out_pdf))
        except Exception:
            return None
        return None

    def _extract_docx_text(self, src: Path) -> str:
        try:
            import mammoth  # type: ignore
            with open(src, 'rb') as f:
                result = mammoth.extract_raw_text(f)
            text = result.value
            if text and text.strip():
                return text
        except Exception:
            pass

        parts: list[str] = []
        try:
            from docx import Document  # type: ignore
            doc = Document(str(src))
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(' | '.join(row_cells))
            if parts:
                return '\n'.join(parts)
        except Exception:
            pass

        try:
            with zipfile.ZipFile(src, 'r') as zf:
                data = zf.read('word/document.xml')
            root = ET.fromstring(data)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            texts = [node.text for node in root.findall('.//w:t', ns) if node.text]
            return '\n'.join(texts)
        except Exception:
            pass
        return ''

    def _extract_doc_text(self, src: Path) -> str:
        text = self._extract_doc_text_via_com(src)
        if text.strip():
            return text

        text = self._extract_doc_text_via_libreoffice(src)
        if text.strip():
            return text

        text = self._extract_doc_text_via_external_tools(src)
        if text.strip():
            return text

        try:
            import textract  # type: ignore
            raw = textract.process(str(src))
            text = raw.decode('utf-8', errors='ignore') if isinstance(raw, bytes) else str(raw)
            text = self._normalize_extracted_text(text)
            if text.strip():
                return text
        except Exception:
            pass

        try:
            import olefile  # type: ignore
            if olefile.isOleFile(str(src)):
                with olefile.OleFileIO(str(src)) as ole:
                    chunks = []
                    for stream_name in ole.listdir():
                        try:
                            data = ole.openstream(stream_name).read()
                        except Exception:
                            continue
                        chunks.extend(self._extract_text_candidates_from_bytes(data))
                    text = self._normalize_extracted_text('\n'.join(chunks))
                    if text.strip():
                        return text
        except Exception:
            pass
        return ''

    def _extract_doc_text_via_com(self, src: Path) -> str:
        try:
            import pythoncom  # type: ignore
            import win32com.client  # type: ignore
        except Exception:
            return ''
        doc = None
        word = None
        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(str(src), ReadOnly=True)
            text = doc.Content.Text
            return text if isinstance(text, str) else ''
        except Exception:
            return ''
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _render_hwp_text_fallback(self, src: Path) -> list[QImage]:
        ext = src.suffix.lower()
        text = ''

        if ext == '.hwpx':
            text = self._extract_hwpx_with_python_hwpx(src) or ''
            if not text:
                try:
                    text = self._extract_hwpx_text(src)
                except Exception:
                    text = ''

        if not text and ext == '.hwp':
            text = self._extract_hwp_with_helper(src) or ''
            if not text:
                text = self._extract_hwp_with_gethwp(src) or ''
            if not text:
                try:
                    import olefile  # type: ignore
                    if olefile.isOleFile(str(src)):
                        with olefile.OleFileIO(str(src)) as ole:
                            chunks = []
                            for stream_name in ole.listdir():
                                try:
                                    data = ole.openstream(stream_name).read()
                                except Exception:
                                    continue
                                chunks.extend(self._extract_text_candidates_from_bytes(data))
                            text = '\n'.join(chunks)
                except Exception:
                    text = text or ''

        if not text:
            return []
        return self._text_to_pages(text, title=src.name)


    def _extract_hwp_with_helper(self, src: Path) -> str:
        try:
            import helper_hwp  # type: ignore
            for attr in ('read_hwp', 'extract_text', 'to_text'):
                fn = getattr(helper_hwp, attr, None)
                if callable(fn):
                    result = fn(str(src))
                    if isinstance(result, str) and result.strip():
                        return result
        except Exception:
            return ''
        return ''

    def _extract_hwp_with_gethwp(self, src: Path) -> str:
        try:
            import gethwp  # type: ignore
            text = gethwp.read_hwp(str(src))
            return text if isinstance(text, str) else ''
        except Exception:
            return ''

    def _extract_hwpx_with_python_hwpx(self, src: Path) -> str:
        try:
            from hwpx import HWPX  # type: ignore
            doc = HWPX(str(src))
            if hasattr(doc, 'get_text'):
                return str(doc.get_text() or '')
        except Exception:
            pass
        try:
            import python_hwpx as hwpx_mod  # type: ignore
            if hasattr(hwpx_mod, 'read_text'):
                return str(hwpx_mod.read_text(str(src)) or '')
        except Exception:
            pass
        return ''

    def _extract_doc_text_via_libreoffice(self, src: Path) -> str:
        soffice = shutil.which('soffice')
        if soffice is None:
            return ''
        out_dir = self._temp_dir / f'libreoffice_txt_{len(self.loaded_documents) + 1}'
        out_dir.mkdir(parents=True, exist_ok=True)
        try:
            result = subprocess.run(
                [soffice, '--headless', '--convert-to', 'txt:Text', '--outdir', str(out_dir), str(src)],
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )
        except Exception:
            return ''
        if result.returncode != 0:
            return ''
        out_txt = out_dir / f'{src.stem}.txt'
        if not out_txt.exists():
            candidates = sorted(out_dir.glob('*.txt'))
            if not candidates:
                return ''
            out_txt = candidates[0]
        try:
            return self._normalize_extracted_text(out_txt.read_text(encoding='utf-8', errors='ignore'))
        except Exception:
            return ''

    def _extract_doc_text_via_external_tools(self, src: Path) -> str:
        for tool in ('antiword', 'catdoc', 'wvText'):
            exe = shutil.which(tool)
            if exe is None:
                continue
            try:
                result = subprocess.run(
                    [exe, str(src)],
                    capture_output=True,
                    text=True,
                    timeout=30,
                    check=False,
                )
            except Exception:
                continue
            output = self._normalize_extracted_text((result.stdout or '') + '\n' + (result.stderr or ''))
            if output.strip():
                return output
        return ''

    def _extract_text_candidates_from_bytes(self, data: bytes) -> list[str]:
        results: list[str] = []
        for encoding in ('utf-16-le', 'utf-8', 'cp949', 'utf-16-be', 'latin1'):
            try:
                decoded = data.decode(encoding, errors='ignore')
            except Exception:
                continue
            cleaned = decoded.replace('\x00', '')
            cleaned = re.sub(r'[\t\r\f\v]+', ' ', cleaned)
            cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
            cleaned = re.sub(r'[^\w가-힣\s\-_,.:;()\[\]/%+*&@!?]', ' ', cleaned)
            cleaned = re.sub(r' {2,}', ' ', cleaned)
            lines = [line.strip() for line in cleaned.splitlines()]
            useful = [line for line in lines if len(re.sub(r'\W+', '', line)) >= 3]
            if useful:
                results.extend(useful[:200])
        uniq: list[str] = []
        seen = set()
        for item in results:
            key = item.strip()
            if key and key not in seen:
                uniq.append(key)
                seen.add(key)
        return uniq

    def _normalize_extracted_text(self, text: str) -> str:
        cleaned = text.replace('\x00', '')
        cleaned = cleaned.replace('\r\n', '\n').replace('\r', '\n')
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
        cleaned = re.sub(r'[ \t]{2,}', ' ', cleaned)
        lines = [line.strip() for line in cleaned.splitlines()]
        useful = [line for line in lines if len(re.sub(r'\W+', '', line)) >= 2]
        return '\n'.join(useful).strip()

    def _build_unreadable_doc_notice(self, src: Path) -> str:
        return '\n'.join([
            'DOC preview fallback',
            '',
            '이 DOC 파일은 현재 환경에서 본문 텍스트를 직접 추출하지 못했습니다.',
            '가능한 경로: Word COM, LibreOffice, antiword/catdoc/wvText, textract, olefile raw fallback.',
            '',
            f'파일명: {src.name}',
            '',
            'Windows 환경에서 정확도를 높이려면:',
            '- Microsoft Word 설치 후 COM 경로 사용',
            '- 또는 LibreOffice 설치 후 headless 변환 경로 사용',
            '- 또는 antiword / catdoc 같은 오픈소스 추출기 설치',
        ])

    def _extract_hwpx_text(self, src: Path) -> str:
        texts: list[str] = []
        with zipfile.ZipFile(src, 'r') as zf:
            xml_names = [n for n in zf.namelist() if n.lower().endswith('.xml')]
            for name in sorted(xml_names):
                if 'section' not in name.lower() and 'contents' not in name.lower() and 'body' not in name.lower():
                    continue
                try:
                    data = zf.read(name)
                    root = ET.fromstring(data)
                except Exception:
                    continue
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        texts.append(elem.text.strip())
        return '\n'.join(texts)

    def _text_to_pages(self, text: str, title: str = '') -> list[QImage]:
        page_size = (1240, 1754)
        margin = 70
        pages: list[QImage] = []
        lines = text.splitlines() or ['']
        batch: list[str] = []
        max_lines = 48
        if title:
            lines = [title, '', *lines]
        for line in lines:
            batch.append(line)
            if len(batch) >= max_lines:
                pages.append(self._render_text_page('\n'.join(batch), page_size, margin))
                batch = []
        if batch:
            pages.append(self._render_text_page('\n'.join(batch), page_size, margin))
        return pages or [self._render_text_page('(빈 문서)', page_size, margin)]

    def _render_text_page(self, text: str, page_size: tuple[int, int], margin: int) -> QImage:
        image = QImage(page_size[0], page_size[1], QImage.Format_ARGB32)
        image.fill(Qt.white)
        painter = QPainter(image)
        doc = QTextDocument()
        font = QFont('Malgun Gothic')
        font.setPointSize(12)
        doc.setDefaultFont(font)
        option = QTextOption()
        option.setWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
        doc.setDefaultTextOption(option)
        plain = text.replace('\r\n', '\n').replace('\r', '\n')
        doc.setPlainText(plain)
        doc.setTextWidth(page_size[0] - margin * 2)
        painter.translate(margin, margin)
        doc.drawContents(painter, QRectF(0, 0, page_size[0] - margin * 2, page_size[1] - margin * 2))
        painter.end()
        return image

    def document_count(self) -> int:
        return len(self.loaded_documents)

    def next_document(self) -> None:
        if not self.loaded_documents:
            return
        self.doc_index = (self.doc_index + 1) % len(self.loaded_documents)
        self.page_index = 0

    def prev_document(self) -> None:
        if not self.loaded_documents:
            return
        self.doc_index = (self.doc_index - 1) % len(self.loaded_documents)
        self.page_index = 0


    def close_current_document(self) -> bool:
        if not self.has_document():
            return False
        current = self.loaded_documents.pop(self.doc_index)
        try:
            if current.doc is not None:
                current.doc.close()
        except Exception:
            pass
        if not self.loaded_documents:
            self.doc_index = -1
            self.page_index = 0
            return True
        self.doc_index = min(self.doc_index, len(self.loaded_documents) - 1)
        self.page_index = 0
        return True

    def page_count(self) -> int:
        current = self.current_document()
        if current is None:
            return 0
        if current.doc is not None:
            return current.doc.page_count
        return len(current.fallback_pages or [])

    def next_page(self) -> None:
        count = self.page_count()
        if count <= 0:
            return
        self.page_index = min(self.page_index + 1, count - 1)

    def prev_page(self) -> None:
        if self.page_count() <= 0:
            return
        self.page_index = max(self.page_index - 1, 0)

    def render_current_page(self, scale: float = 2.0) -> QImage | None:
        current = self.current_document()
        if current is None:
            return None
        if current.doc is not None:
            page = current.doc.load_page(self.page_index)
            mat = fitz.Matrix(scale, scale)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            fmt = QImage.Format_RGB888
            image = QImage(pix.samples, pix.width, pix.height, pix.stride, fmt).copy()
            return image
        pages = current.fallback_pages or []
        if not pages:
            return None
        return pages[self.page_index].copy()

    def render_current_clip(self, image_rect: QRectF, base_render_scale: float = 2.0, output_scale: float = 3.0) -> QImage | None:
        current = self.current_document()
        if current is None:
            return None
        if current.doc is None:
            page = self.render_current_page(scale=base_render_scale)
            if page is None:
                return None
            return page.copy(int(image_rect.x()), int(image_rect.y()), int(image_rect.width()), int(image_rect.height()))

        page = current.doc.load_page(self.page_index)
        scale_ratio = output_scale / max(base_render_scale, 1e-6)
        clip = fitz.Rect(
            image_rect.x() / base_render_scale,
            image_rect.y() / base_render_scale,
            (image_rect.x() + image_rect.width()) / base_render_scale,
            (image_rect.y() + image_rect.height()) / base_render_scale,
        )
        pix = page.get_pixmap(matrix=fitz.Matrix(output_scale, output_scale), clip=clip, alpha=False)
        fmt = QImage.Format_RGB888
        image = QImage(pix.samples, pix.width, pix.height, pix.stride, fmt).copy()
        return image
